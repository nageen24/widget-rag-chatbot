"""
Website scraper for hewmannexperience.com
Single-page WordPress/Elementor site — extracts all sections into a Word document.
"""

import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

BASE_URL = "https://hewmannexperience.com"

SECTION_IDS = ["content", "services", "about", "pricing", "faqs"]


def clean_text(text):
    """Clean up whitespace from text."""
    return " ".join(text.split()).strip()


def extract_section(section_el):
    """Extract all meaningful text from a section element."""
    results = []
    seen = set()

    # Walk through all elements in document order
    for el in section_el.find_all(True):
        tag = el.name

        # Skip non-content tags
        if tag in ["script", "style", "noscript", "svg", "path", "link", "meta"]:
            continue

        # Only process leaf-ish content nodes
        if tag in ["h1", "h2", "h3", "h4", "h5", "h6", "p", "li", "td", "th", "label", "span", "a"]:
            # Avoid spans/divs that just wrap other block elements
            if tag in ["span", "a"] and el.find(["h1","h2","h3","h4","p","li"]):
                continue
            text = clean_text(el.get_text(separator=" "))
            if text and text not in seen and len(text) > 10:
                seen.add(text)
                results.append((tag, text))

    return results


def scrape():
    print(f"Fetching {BASE_URL} ...")
    resp = requests.get(BASE_URL, timeout=20, headers={
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"
    })
    resp.raise_for_status()
    soup = BeautifulSoup(resp.text, "lxml")

    # Also grab page title and meta description
    page_title = soup.title.get_text(strip=True) if soup.title else "Hewmann Experience"
    meta_desc = ""
    meta = soup.find("meta", attrs={"name": "description"})
    if meta:
        meta_desc = meta.get("content", "")

    print(f"  Page title: {page_title}")
    print(f"  Meta description: {meta_desc}")

    pages_data = []

    # --- Hero / full page content first ---
    hero = soup.find(id="content") or soup.find("main")
    if hero:
        pages_data.append({
            "section": "Homepage Overview",
            "content": extract_section(hero)
        })

    # --- Named sections ---
    for sec_id in ["services", "about", "pricing", "faqs"]:
        el = soup.find(id=sec_id)
        if el:
            print(f"  Found section: #{sec_id}")
            content = extract_section(el)
            pages_data.append({
                "section": sec_id.capitalize(),
                "content": content
            })
        else:
            print(f"  Section #{sec_id} not found")

    # --- Footer ---
    footer = soup.find("footer") or soup.find(id="colophon")
    if footer:
        content = extract_section(footer)
        if content:
            pages_data.append({
                "section": "Contact & Footer",
                "content": content
            })

    return page_title, meta_desc, pages_data


def build_word_doc(page_title, meta_desc, pages_data, output_path):
    doc = Document()

    # Cover heading
    h = doc.add_heading(page_title, 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if meta_desc:
        p = doc.add_paragraph(meta_desc)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Source: {BASE_URL}")
    doc.add_page_break()

    heading_tags = {"h1", "h2", "h3", "h4", "h5", "h6"}

    for page in pages_data:
        doc.add_heading(page["section"], level=1)
        seen = set()

        for tag, text in page["content"]:
            if text in seen:
                continue
            seen.add(text)

            if tag in heading_tags:
                level = int(tag[1]) if len(tag) == 2 and tag[1].isdigit() else 2
                doc.add_heading(text, level=min(level, 4))
            elif tag == "li":
                doc.add_paragraph(text, style="List Bullet")
            else:
                doc.add_paragraph(text)

        doc.add_page_break()

    doc.save(output_path)
    print(f"\nSaved Word doc: {output_path}")


if __name__ == "__main__":
    page_title, meta_desc, pages_data = scrape()

    print(f"\nExtracted {len(pages_data)} sections.")

    os.makedirs("d:/EWD Projects/rag-chatbot/data", exist_ok=True)
    output = "d:/EWD Projects/rag-chatbot/data/hewmann_experience_knowledge_base.docx"
    build_word_doc(page_title, meta_desc, pages_data, output)

    print("\nSections extracted:")
    for p in pages_data:
        print(f"  - {p['section']}: {len(p['content'])} items")
